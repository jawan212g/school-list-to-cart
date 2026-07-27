"""Deterministic extraction and matching support for the offline demo path."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from agent.extract import ExtractionInputError
from agent.normalize import canonicalize_item_name
from agent.schema import ExtractionEnvelope, Requirement


DEMO_LIST_TEXT = """\
12 #2 pencils
4 glue sticks
2 composition notebooks
1 box of tissues
1 pair of headphones
Optional: 1 backpack
"""

_QUANTITY_PATTERN = re.compile(r"^\s*(\d+)\s+(.*)$")
_UNIT_WORDS = frozenset(
    {"box", "boxes", "pack", "packs", "pair", "pairs", "of"}
)


def _read_demo_text(
    source: str | Path | bytes,
    mime_type: str | None,
) -> str:
    """Read text-like demo input without a network or model call."""

    if isinstance(source, bytes):
        if mime_type == "text/plain":
            return source.decode("utf-8")
        if mime_type == (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ):
            document = Document(BytesIO(source))
            blocks = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]
            for table in document.tables:
                for row in table.rows:
                    cells = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if cells:
                        blocks.append(" | ".join(cells))
            return "\n".join(blocks)
        if mime_type == "application/pdf":
            reader = PdfReader(BytesIO(source))
            return "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        raise ExtractionInputError(
            "Offline demo mode supports pasted text, TXT, DOCX, PDF, "
            "and the bundled sample. Use normal mode for arbitrary images."
        )
    if isinstance(source, Path):
        return _read_demo_text(source.read_bytes(), mime_type)
    return source


def _canonical_from_text(text: str) -> str | None:
    words = [
        word
        for word in re.sub(r"[^a-z0-9#]+", " ", text.casefold()).split()
        if word not in _UNIT_WORDS
    ]
    for start in range(len(words)):
        for end in range(len(words), start, -1):
            canonical = canonicalize_item_name("_".join(words[start:end]))
            if canonical is not None:
                return canonical
    return None


def extract_demo_document(
    source: str | Path | bytes,
    *,
    child_id: str = "unassigned",
    mime_type: str | None = None,
    client: object | None = None,
) -> ExtractionEnvelope:
    """Extract a simple demo list deterministically without external services."""

    del client
    text = _read_demo_text(source, mime_type)
    requirements: list[Requirement] = []
    for index, raw_line in enumerate(text.splitlines(), start=1):
        line = raw_line.strip(" \t•-")
        if not line:
            continue
        optional = line.casefold().startswith("optional:")
        if optional:
            line = line.split(":", 1)[1].strip()
        quantity_match = _QUANTITY_PATTERN.match(line)
        if quantity_match is None:
            quantity = 1
            description = line
            confidence = 0.6
        else:
            quantity = int(quantity_match.group(1))
            description = quantity_match.group(2)
            confidence = 1.0
        canonical = _canonical_from_text(description)
        if canonical is None:
            continue
        brand_lock = (
            "Ticonderoga"
            if "ticonderoga" in description.casefold()
            else None
        )
        requirements.append(
            Requirement(
                req_id=f"demo-{index}",
                child_id=child_id,
                raw_text=raw_line.strip(),
                canonical_item=canonical,
                quantity=quantity,
                unit_type="each",
                brand_lock=brand_lock,
                is_required=not optional,
                is_purchasable=True,
                requirement_type=(
                    "optional" if optional else "required"
                ),
                extraction_confidence=confidence,
            )
        )
    if not requirements:
        raise ExtractionInputError(
            "Offline demo mode could not identify readable school-supply items"
        )
    return ExtractionEnvelope(requirements=tuple(requirements))
