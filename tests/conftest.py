"""Repeatable file fixtures for supply-list intake tests."""

from __future__ import annotations

import json
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from agent.match import (
    SuitabilityCase,
    SuitabilityDecision,
)
from agent.schema import ExtractionEnvelope


@dataclass(frozen=True)
class FrozenSuitabilityJudge:
    """Replay the captured model decisions for production match cases."""

    decisions: Mapping[tuple[str, str], SuitabilityDecision]

    def judge(
        self,
        cases: Sequence[SuitabilityCase],
    ) -> tuple[SuitabilityDecision, ...]:
        """Return one frozen decision for every current production case."""

        missing = tuple(
            (case.need_key, case.offer.sku)
            for case in cases
            if (case.need_key, case.offer.sku) not in self.decisions
        )
        if missing:
            raise AssertionError(
                "Frozen Maple suitability fixture is missing current cases: "
                f"{missing[:3]}"
            )
        return tuple(
            self.decisions[(case.need_key, case.offer.sku)]
            for case in cases
        )


@dataclass(frozen=True)
class FrozenMapleFixture:
    """Confirmed Maple extractions plus captured suitability decisions."""

    metadata: Mapping[str, object]
    extractions: Mapping[str, ExtractionEnvelope]
    judge: FrozenSuitabilityJudge


def load_frozen_maple_fixture() -> FrozenMapleFixture:
    """Load the production-shaped Maple boundary outside pytest injection."""

    fixture_path = (
        Path(__file__).parent
        / "fixtures"
        / "maple_street_frozen_pipeline.json"
    )
    payload = json.loads(fixture_path.read_text(encoding="utf-8"))
    decisions = tuple(
        SuitabilityDecision.model_validate(decision)
        for decision in payload["suitability_decisions"]
    )
    return FrozenMapleFixture(
        metadata=payload["fixture_metadata"],
        extractions={
            child_id: ExtractionEnvelope.model_validate(envelope)
            for child_id, envelope in payload["envelopes"].items()
        },
        judge=FrozenSuitabilityJudge(
            decisions={
                (decision.need_key, decision.sku): decision
                for decision in decisions
            }
        ),
    )


@pytest.fixture(scope="session")
def frozen_maple_fixture() -> FrozenMapleFixture:
    """Load the production-shaped, model-free Maple cart regression input."""

    return load_frozen_maple_fixture()


@pytest.fixture
def docx_list_bytes() -> bytes:
    """Return a deterministic DOCX with paragraphs, a bullet, and a table."""

    document = Document()
    document.add_paragraph("2 boxes of tissues")
    document.add_paragraph("12 pencils", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "4 glue sticks"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _image_bytes(image_format: str) -> bytes:
    image = Image.new("RGB", (16, 16), color=(255, 255, 255))
    output = BytesIO()
    image.save(output, format=image_format)
    return output.getvalue()


@pytest.fixture
def png_list_bytes() -> bytes:
    """Return stable PNG bytes for image-content packaging tests."""

    return _image_bytes("PNG")


@pytest.fixture
def jpeg_list_bytes() -> bytes:
    """Return stable JPEG bytes for image-content packaging tests."""

    return _image_bytes("JPEG")


@pytest.fixture
def multipage_pdf_bytes() -> bytes:
    """Return a two-page PDF whose visual page boundaries are deterministic."""

    first = Image.new("RGB", (120, 80), color=(255, 255, 255))
    second = Image.new("RGB", (120, 80), color=(245, 245, 245))
    output = BytesIO()
    first.save(
        output,
        format="PDF",
        save_all=True,
        append_images=(second,),
    )
    return output.getvalue()
